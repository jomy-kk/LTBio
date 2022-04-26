
from collections import OrderedDict
import docx as dc
import datetime
import json
import os
import pandas as pd
import string

import sys
sys.path.append('C:\\Users\\Mariana\\PycharmProjects\\mapme\\code')
import read_files.read_hospital as rh

# class Report_HEM:
#     def __init__(self, patient_id: str):
#
#         assert patient_id in ['358', '400', '413', '352', '396', '363']
#         self.id = patient_id
#
#     def report(self):activ


def simplify_class(original):
    new_ilae = []
    inp_types = ['FAS', 'FWIA', 'FBTC', 'FUnk', 'None']
    for seiz in original:
        inp = ''
        while inp not in inp_types:
            print(f'Choose one from {inp_types}')
            inp = input(f'{seiz} from {original}')
            yn = input(f'Want to save {seiz} as {inp} y/n')
            if yn.lower() != 'y':
                inp = ''
        new_ilae += [inp]
    return new_ilae


def report_358(dir):

    doc_dir = [file for file in os.listdir(dir) if file.endswith('.docx')][0]

    # get word document
    doc = dc.Document(os.path.join(dir, doc_dir))

    # find table that has 'V/S', which is the table with the dates
    try:
        doc_table = [table for table in doc._body.tables if [table for cell in table._cells if 'V/S' in cell.text]!=[]][0]
    except:
        print('V/S is not in tables, please check your report.')
        doc_table = None


    seizure_times = pd.DataFrame()

    seizure_times_type, seizure_times_date, seizure_times_state, seizure_times_class = [], [], [], []

    for cl, cell in enumerate(doc_table._cells):

        if ('Evento' in cell.text) or ('Crise' in cell.text):
            try:
                date = datetime.datetime.strptime(doc_table._cells[cl + 1].text, '%H:%M:%S\n%d-%m-%Y')
            except:
                date = None
            if date:
                seizure_times_type += [cell.text]
                seizure_times_date += [datetime.datetime.strftime(date,'%d-%m-%Y\n%H:%M:%S')]
                seizure_times_state += [doc_table._cells[cl + 2].text]
                if doc_table._cells[cl + 3].text == doc_table._cells[cl + 2].text:
                    idx = cl + 4
                else:
                    idx = cl + 3
                seizure_times_class += [doc_table._cells[idx].text]

    seizure_times['Type'] = seizure_times_type
    seizure_times['Date'] = seizure_times_date
    seizure_times['Class'] = seizure_times_class
    seizure_times['State'] = seizure_times_state
    seizure_times['ILAE'] = simplify_class(seizure_times_class)

    seizure_times.to_csv(os.path.join(dir, 'seizure_label'), columns = seizure_times.columns)
    return seizure_times


def report_400(dir):

    doc_dir = [file for file in os.listdir(dir) if file.endswith('.docx')][0]

    # get word document
    doc = dc.Document(os.path.join(dir, doc_dir))

    # find table that has 'V/S', which is the table with the dates
    try:
        doc_table = \
        [table for table in doc._body.tables if [table for cell in table._cells if 'V/S' in cell.text] != []][0]
    except:
        print('V/S is not in tables, please check your report.')
        doc_table = None

    seizure_times = pd.DataFrame()

    seizure_times_type, seizure_times_date, seizure_times_state, seizure_times_class = [], [], [], []

    for cl, cell in enumerate(doc_table._cells):

        if cell.text.strip() in string.digits or cell.text.strip() in ['8-11']:
            print(cell.text)
            if '-' in cell.text:
                date_strings = doc_table._cells[cl + 1].text.split('\n')
                date = date_strings[-1]
                for times in date_strings[:-2]:
                    seizure_times_date += [date + '\n' + times]
                    seizure_times_class += [doc_table._cells[cl + 3].text.strip()]
                    seizure_times_state += [doc_table._cells[cl + 2].text.strip()]
                    seizure_times_type += [cell.text.strip()]

            try:
                date = datetime.datetime.strptime(doc_table._cells[cl + 1].text, '%H:%M:%S\n%d-%m-%Y')
            except:
                date = None
            if date:
                seizure_times_type += [cell.text]
                seizure_times_date += [datetime.datetime.strftime(date,'%d-%m-%Y\n%H:%M:%S')]
                seizure_times_state += [doc_table._cells[cl + 2].text]
                if doc_table._cells[cl + 3].text == doc_table._cells[cl + 2].text:
                    idx = cl + 4
                else:
                    idx = cl + 3
                seizure_times_class += [doc_table._cells[idx].text]


    seizure_times['Type'] = seizure_times_type
    seizure_times['Date'] = seizure_times_date
    seizure_times['Class'] = seizure_times_class
    seizure_times['State'] = seizure_times_state
    seizure_times['ILAE'] = simplify_class(seizure_times_class)

    seizure_times.to_csv(os.path.join(dir, 'seizure_label'), columns=seizure_times.columns)
    return seizure_times

#report_400(dir = 'F:\\Patients_HEM\\Retrospective\\PAT_400')

def report_326(dir):

    doc_dir = [file for file in os.listdir(dir) if file.endswith('.docx')][0]


    # get word document
    doc = dc.Document(os.path.join(dir, doc_dir))

    # find table that has 'V/S', which is the table with the dates
    try:
        doc_table = \
        [table for table in doc._body.tables if [table for cell in table._cells if 'V/S' in cell.text] != []][0]
    except:
        print('V/S is not in tables, please check your report.')
        doc_table = None

    seizure_times = pd.DataFrame()

    seizure_times_type, seizure_times_date, seizure_times_state, seizure_times_class = [], [], [], []

    for cl, cell in enumerate(doc_table._cells):

        if cell.text.strip() in string.digits or '1' in cell.text:

            if '-' in cell.text:
                date_ = doc_table._cells[cl + 1].text.split('\n')[-1]

            try:
                date = datetime.datetime.strptime(doc_table._cells[cl + 1].text.strip(), '%d-%m-%Y\n%H:%M:%S')
            except:
                date = None

            if date:
                seizure_times_type += [cell.text.strip()]
                seizure_times_date += [datetime.datetime.strftime(date,'%d-%m-%Y\n%H:%M:%S')]
                seizure_times_state += [doc_table._cells[cl + 2].text.strip()]
                if doc_table._cells[cl + 3].text == doc_table._cells[cl + 2].text:
                    idx = cl + 4
                else:
                    idx = cl + 3
                seizure_times_class += [doc_table._cells[idx].text.strip()]
            print(cell.text, date, )


    seizure_times['Type'] = seizure_times_type
    seizure_times['Date'] = seizure_times_date
    seizure_times['Class'] = seizure_times_class
    seizure_times['State'] = seizure_times_state
    seizure_times['ILAE'] = simplify_class(seizure_times_class)

    seizure_times.to_csv(os.path.join(dir, 'seizure_label'), columns=seizure_times.columns)
    return seizure_times

#report_326('F:\\Patients_HEM\\Retrospective\\PAT_326_EXAMES')

def report_413(dir):

    doc_dir = [file for file in os.listdir(dir) if file.endswith('.docx')][0]

    # get word document
    doc = dc.Document(os.path.join(dir, doc_dir))

    # find table that has 'V/S', which is the table with the dates
    try:
        doc_table = \
        [table for table in doc._body.tables if [table for cell in table._cells if 'V/S' in cell.text] != []][0]
    except:
        print('V/S is not in tables, please check your report.')
        doc_table = None

    seizure_times = pd.DataFrame()

    seizure_times_type, seizure_times_date, seizure_times_state, seizure_times_class = [], [], [], []

    for cl, cell in enumerate(doc_table._cells):


        if cell.text.strip() in string.digits:
            print(cell.text)
            if '-' in cell.text:
                date_ = doc_table._cells[cl + 1].text.split('\n')[-1]


            try:
                date = datetime.datetime.strptime(doc_table._cells[cl + 1].text.strip(), '%d/%m/%Y\n%H:%M:%S')
            except:
                date = None
            if date:
                seizure_times_type += [cell.text.strip()]
                seizure_times_date += [datetime.datetime.strftime(date,'%d-%m-%Y\n%H:%M:%S')]
                seizure_times_state += [doc_table._cells[cl + 2].text.strip()]
                if doc_table._cells[cl + 3].text == doc_table._cells[cl + 2].text:
                    idx = cl + 4
                else:
                    idx = cl + 3
                seizure_times_class += [doc_table._cells[idx].text.strip()]

    seizure_times['Type'] = seizure_times_type
    seizure_times['Date'] = seizure_times_date
    seizure_times['Class'] = seizure_times_class
    seizure_times['State'] = seizure_times_state
    seizure_times['ILAE'] = simplify_class(seizure_times_class)

    seizure_times.to_csv(os.path.join(dir, 'seizure_label'), columns=seizure_times.columns)
    return seizure_times


def report_352(dir):

    doc_dir = [file for file in os.listdir(dir) if file.endswith('.docx')][0]

    # get word document
    doc = dc.Document(os.path.join(dir, doc_dir))

    # find table that has 'V/S', which is the table with the dates
    try:
        doc_table = \
        [table for table in doc._body.tables if [table for cell in table._cells if 'V/S' in cell.text] != []][0]
    except:
        print('V/S is not in tables, please check your report.')
        doc_table = None

    seizure_times = pd.DataFrame()

    seizure_times_type, seizure_times_date, seizure_times_state, seizure_times_class = [], [], [], []

    for cl, cell in enumerate(doc_table._cells):
        if '1' in cell.text:

            try:
                date = datetime.datetime.strptime(doc_table._cells[cl + 1].text, '\n%d-%m-%Y\n%H:%M:%S')
            except:
                date = None
            if date:
                seizure_times_type += [cell.text]
                seizure_times_date += [datetime.datetime.strftime(date,'%d-%m-%Y\n%H:%M:%S')]
                seizure_times_state += [doc_table._cells[cl + 2].text]
                if doc_table._cells[cl + 3].text == doc_table._cells[cl + 2].text:
                    idx = cl + 4
                else:
                    idx = cl + 3
                seizure_times_class += [doc_table._cells[idx].text]

    seizure_times['Type'] = seizure_times_type
    seizure_times['Date'] = seizure_times_date
    seizure_times['Class'] = seizure_times_class
    seizure_times['State'] = seizure_times_state
    seizure_times['ILAE'] = simplify_class(seizure_times_class)

    seizure_times.to_csv(os.path.join(dir, 'seizure_label'), columns=seizure_times.columns)
    return seizure_times


def report_386(dir):

    doc_dir = [file for file in os.listdir(dir) if file.endswith('.docx')][0]

    # get word document
    doc = dc.Document(os.path.join(dir, doc_dir))

    # find table that has 'V/S', which is the table with the dates
    try:
        doc_table = \
            [table for table in doc._body.tables if [table for cell in table._cells if 'V/S' in cell.text] != []][0]
    except:
        print('V/S is not in tables, please check your report.')
        doc_table = None

    seizure_times = pd.DataFrame()

    seizure_times_type, seizure_times_date, seizure_times_state, seizure_times_class = [], [], [], []

    for cl, cell in enumerate(doc_table._cells):

        if cell.text.strip() in string.digits:

            try:
                date = datetime.datetime.strptime(doc_table._cells[cl + 1].text.strip(), '%H:%M:%S\n%d-%m-%Y')
            except:
                date = None
            if date:
                seizure_times_type += [cell.text.strip()]
                seizure_times_date += [datetime.datetime.strftime(date, '%d-%m-%Y\n%H:%M:%S')]
                seizure_times_state += [doc_table._cells[cl + 2].text]
                if doc_table._cells[cl + 3].text == doc_table._cells[cl + 2].text:
                    idx = cl + 4
                else:
                    idx = cl + 3
                seizure_times_class += [doc_table._cells[idx].text.strip()]

    seizure_times['Type'] = seizure_times_type
    seizure_times['Date'] = seizure_times_date
    seizure_times['Class'] = seizure_times_class
    seizure_times['State'] = seizure_times_state
    seizure_times['ILAE'] = simplify_class(seizure_times_class)

    seizure_times.to_csv(os.path.join(dir, 'seizure_label'), columns=seizure_times.columns)
    return seizure_times


def read_excel_report(pat, excel_name='Patients_HSM_.xlsx', dir='E:\\Patients_HSM'):

    report = pd.read_excel(os.path.join(dir, excel_name), sheet_name = pat, header=0)
    seizure_times = pd.DataFrame()
    seizure_times_type, seizure_times_date, seizure_times_class, seizure_times_state = [], [], [], []
    seizure_times_loc, seizure_times_side = [], []

    for row in range(len(report)):
        try:
            date = datetime.datetime.strptime(report.iloc[row]['Data'], '%d-%m-%Y') \
                if type(report.iloc[row]['Data']) == str else report.iloc[row]['Data']
            hour = datetime.datetime.strptime(report.iloc[row]['Hora Clínica'], '%H:%M:%S') \
                if type(report.iloc[row]['Hora Clínica']) == str else report.iloc[row]['Hora Clínica']
            date = datetime.datetime.combine(date, hour)
        except:
            date = None
        print(date)
        if date:
            seizure_times_type += [report.iloc[row]['Crises']]
            seizure_times_date += [date]
            seizure_times_class += [report.iloc[row]['Focal / Generalisada']]
            seizure_times_state += [report.iloc[row]['Sono/ Vigília']]
            seizure_times_loc += [report.iloc[row]['Localização']]
            seizure_times_side += [report.iloc[row]['lado']]

    seizure_times['Type'] = seizure_times_type
    seizure_times['Date'] = seizure_times_date
    seizure_times['Class'] = seizure_times_class
    seizure_times['State'] = seizure_times_state
    seizure_times['Loc'] = seizure_times_loc
    seizure_times['Side'] = seizure_times_side
    seizure_times.to_csv(os.path.join(dir, pat, 'seizure_label'), columns=seizure_times.columns)

    return seizure_times



if __name__ == '__main__':

    dir = 'F:\\PreEpiSeizures\\Patients_HEM\\Retrospective'
    for patient in sorted(os.listdir(dir)):
        pat_dir = os.path.join(dir, patient)
        if not os.path.isdir(os.path.join(pat_dir)):
            continue
        if patient in ['PAT_312_EXAMES', 'PAT_326_EXAMES', 'PAT_352_EXAMES', 'PAT_358', 'PAT_365_EXAMES']:
            continue
        else:
            if 'PAT' in patient:
                id = patient.split('_')[1]

            if id == '326':
                report_326(pat_dir)
            elif id == '352':
                report_352(pat_dir)
            elif id == '358':
                report_358(pat_dir)
            elif id == '386':
                report_386(pat_dir)
            elif id == '400':
                report_400(pat_dir)
            elif id == '413':
                report_413(pat_dir)
            else:
                print('Patient ' + patient + ' does not have read report function')
                rh.read_trc_events(os.path.join(pat_dir, 'hospital'), pat_dir)
            print(f'{patient}')

# for pat in os.listdir('E:\\Patients_HSM'):
#
#     try:
#         read_excel_report(pat)
#         print('success')
#     except:
#         continue
#report_400()
#json.dumps(seizure_times, open(os.path.join(dir,'seizure_times.json'),'w'), sort_keys=True, indent=4)

